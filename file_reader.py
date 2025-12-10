import os
import re
from typing import Optional, List, Tuple
import tempfile

class FileReader:
    """Класс для чтения различных форматов файлов"""
    
    def __init__(self):
        self.supported_formats = ['.txt', '.pdf', '.doc', '.docx']
    
    def read_file(self, file_path: str) -> Optional[str]:
        """
        Чтение файла любого поддерживаемого формата
        
        Args:
            file_path: путь к файлу
            
        Returns:
            текст из файла или None в случае ошибки
        """
        try:
            ext = os.path.splitext(file_path)[1].lower()
            
            if ext == '.txt':
                return self._read_txt(file_path)
            elif ext == '.pdf':
                return self._read_pdf(file_path)
            elif ext == '.doc':
                return self._read_doc(file_path)
            elif ext == '.docx':
                return self._read_docx(file_path)
            else:
                raise ValueError(f"Неподдерживаемый формат файла: {ext}")
                
        except Exception as e:
            print(f"Ошибка чтения файла {file_path}: {e}")
            return None
    
    def _read_txt(self, file_path: str) -> str:
        """Чтение текстового файла"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def _read_pdf(self, file_path: str) -> str:
        """Чтение PDF файла"""
        try:
            import PyPDF2
            return self._read_pdf_pypdf2(file_path)
        except ImportError:
            try:
                import pdfplumber
                return self._read_pdf_pdfplumber(file_path)
            except ImportError:
                raise ImportError(
                    "Для чтения PDF файлов необходимо установить PyPDF2 или pdfplumber. "
                    "Установите: pip install PyPDF2 pdfplumber"
                )
    
    def _read_pdf_pypdf2(self, file_path: str) -> str:
        """Чтение PDF с использованием PyPDF2"""
        import PyPDF2
        
        text = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        
        return '\n'.join(text)
    
    def _read_pdf_pdfplumber(self, file_path: str) -> str:
        """Чтение PDF с использованием pdfplumber"""
        import pdfplumber
        
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        
        return '\n'.join(text)
    
    def _read_doc(self, file_path: str) -> str:
        """Чтение старого формата DOC"""
        try:
            # Попытка 1: через antiword (если установлен в системе)
            return self._read_doc_antiword(file_path)
        except:
            # Попытка 2: через catdoc
            try:
                return self._read_doc_catdoc(file_path)
            except:
                # Попытка 3: через python-docx (может работать с некоторыми .doc)
                try:
                    return self._read_docx(file_path)
                except:
                    raise Exception(
                        "Для чтения .doc файлов требуется установка дополнительных инструментов.\n"
                        "Для Linux: установите antiword или catdoc\n"
                        "Для Windows: конвертируйте файл в .docx или .txt"
                    )
    
    def _read_doc_antiword(self, file_path: str) -> str:
        """Чтение DOC через antiword (Linux/Mac)"""
        import subprocess
        
        try:
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            return result.stdout
        except FileNotFoundError:
            raise Exception(
                "Для чтения .doc файлов требуется установить antiword.\n"
                "Установите: sudo apt-get install antiword (Linux) или brew install antiword (Mac)"
            )
    
    def _read_doc_catdoc(self, file_path: str) -> str:
        """Чтение DOC через catdoc (Linux/Mac)"""
        import subprocess
        
        try:
            result = subprocess.run(
                ['catdoc', '-w', file_path],
                capture_output=True,
                text=True,
                encoding='utf-8',
                errors='ignore'
            )
            return result.stdout
        except FileNotFoundError:
            raise Exception(
                "Для чтения .doc файлов требуется установить catdoc.\n"
                "Установите: sudo apt-get install catdoc (Linux) или brew install catdoc (Mac)"
            )
    
    def _read_docx(self, file_path: str) -> str:
        """Чтение DOCX файла"""
        try:
            import docx
            return self._read_docx_python(file_path)
        except ImportError:
            raise ImportError(
                "Для чтения DOCX файлов необходимо установить python-docx. "
                "Установите: pip install python-docx"
            )
    
    def _read_docx_python(self, file_path: str) -> str:
        """Чтение DOCX с использованием python-docx"""
        import docx
        
        doc = docx.Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text.append(' '.join(row_text))
        
        return '\n'.join(text)
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Получение информации о файле
        
        Args:
            file_path: путь к файлу
            
        Returns:
            словарь с информацией о файле
        """
        try:
            ext = os.path.splitext(file_path)[1].lower()
            file_size = os.path.getsize(file_path)
            file_size_kb = file_size / 1024
            
            # Получаем имя файла без расширения
            filename = os.path.basename(file_path)
            filename_without_ext = os.path.splitext(filename)[0]
            
            # Читаем текст для подсчета статистики
            text = self.read_file(file_path)
            text_length = len(text) if text else 0
            words_count = len(text.split()) if text else 0
            
            return {
                'extension': ext,
                'size_kb': round(file_size_kb, 2),
                'size_bytes': file_size,
                'filename': filename,
                'filename_without_ext': filename_without_ext,
                'text_length': text_length,
                'words_count': words_count
            }
            
        except Exception as e:
            return {
                'error': str(e),
                'extension': os.path.splitext(file_path)[1].lower(),
                'size_kb': os.path.getsize(file_path) / 1024
            }
    
    def convert_to_txt(self, file_path: str, output_path: Optional[str] = None) -> str:
        """
        Конвертация файла в текстовый формат
        
        Args:
            file_path: путь к исходному файлу
            output_path: путь для сохранения (если None - создается временный файл)
            
        Returns:
            путь к созданному текстовому файлу
        """
        text = self.read_file(file_path)
        if not text:
            raise Exception(f"Не удалось прочитать файл: {file_path}")
        
        if output_path is None:
            # Создаем временный файл
            fd, temp_path = tempfile.mkstemp(suffix='.txt')
            os.close(fd)
            with open(temp_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return temp_path
        else:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return output_path